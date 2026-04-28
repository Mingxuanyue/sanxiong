# -*- coding: utf-8 -*-
"""fix_docs.py — 批量修复设计文档v2和用户文档v2中的错误"""
try:
    from docx import Document
    from docx.shared import Pt
    import os, copy

    BASE = r'd:\Users\ymx36\Desktop\三雄争锋游戏'

    # ── 通用：段落文字替换 ────────────────────────────────────────────────
    def replace_in_para(para, old, new):
        """在段落中替换文字，保留原有格式 run"""
        if old not in para.text:
            return False
        # 逐 run 替换
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
        # 如果 runs 拼接才出现（跨run），退而全段文字替换（简单处理）
        if old in para.text:
            full = para.text
            para.clear()
            para.add_run(full.replace(old, new))
        return True

    def replace_in_doc(doc, replacements):
        count = 0
        for para in doc.paragraphs:
            for old, new in replacements.items():
                if replace_in_para(para, old, new):
                    count += 1
        # 表格内
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old, new in replacements.items():
                            if replace_in_para(para, old, new):
                                count += 1
        return count

    # ══════════════════════════════════════════════════════════════════════
    # 1. 设计文档 v2
    # ══════════════════════════════════════════════════════════════════════
    path_design = os.path.join(BASE, '游戏设计文档v2.docx')
    doc = Document(path_design)

    design_replacements = {
        '大獋': '大王',
        '小獋': '小王',
        '大獋防御': '大王防御',
        '小獋偷窥': '小王偷窥',
        '大小獋同台': '大小王同台',
    }
    n = replace_in_doc(doc, design_replacements)
    print(f"设计文档：替换 {n} 处")

    # 在文档开头段落后插入小组名称 + 下载链接（第1段通常是标题）
    # 找第一个非空段落
    insert_idx = 0
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            insert_idx = i + 1
            break

    # 插入小组名称段落（在标题后）
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def insert_paragraph_after(doc, ref_para, text, bold=False, size=12):
        new_para = OxmlElement('w:p')
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size * 2))
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        new_para.append(r)
        ref_para._element.addnext(new_para)
        return new_para

    # 找到标题段（含"三雄争锋"的段落）
    title_para = None
    for p in doc.paragraphs:
        if '三雄争锋' in p.text and '游戏设计文档' in p.text:
            title_para = p
            break
    if title_para:
        insert_paragraph_after(doc, title_para,
            '开发团队：不知道叫啥小组', bold=True, size=12)
        insert_paragraph_after(doc, title_para,
            '下载地址：github.com/Mingxuanyue/sanxiong/releases', bold=False, size=11)

    doc.save(path_design)
    print(f"设计文档已保存：{path_design}")

    # ══════════════════════════════════════════════════════════════════════
    # 2. 用户文档 v2
    # ══════════════════════════════════════════════════════════════════════
    path_user = os.path.join(BASE, '游戏用户文档v2.docx')
    doc2 = Document(path_user)

    user_replacements = {
        '大獋': '大王',
        '小獋': '小王',
        '大獋防御': '大王防御',
        '小獋偷窥': '小王偷窥',
        '大小獋同台': '大小王同台',
        '续尺时': '续场时',
        '比大/比小见则继承': '比大/比小规则继承',
    }
    n2 = replace_in_doc(doc2, user_replacements)
    print(f"用户文档：替换 {n2} 处")

    # 修复 2.1 运行游戏：在"编译后双击运行"段落后加下载方式
    for i, p in enumerate(doc2.paragraphs):
        if '编译后双击运行' in p.text or '编译环境' in p.text:
            # 找到后在下方插入下载段落
            insert_paragraph_after(doc2, p,
                '【推荐】直接下载：无需编译，前往 github.com/Mingxuanyue/sanxiong/releases 下载 sanxiong.exe，双击运行即可，无需安装任何依赖。',
                bold=False, size=11)
            print("用户文档：已插入下载方式说明")
            break

    # 检测并删除重复章节（第四章速查表在文档中出现两次）
    # 找所有含"4.1 基础积分速查"的段落
    positions = []
    for i, p in enumerate(doc2.paragraphs):
        if '4.1' in p.text and '基础积分速查' in p.text:
            positions.append(i)

    if len(positions) >= 2:
        # 删除第二次出现的第四章及之后直到文档末尾的重复内容
        # 找第五章在哪里第二次出现
        fifth_positions = []
        for i, p in enumerate(doc2.paragraphs):
            if '第五章' in p.text and '常见问题' in p.text:
                fifth_positions.append(i)

        if len(fifth_positions) >= 2:
            # 从第二个"第四章"开始的段落全部删除（重复内容）
            second_ch4_para_idx = positions[1]
            # 找到对应段落的 XML element
            paras_to_remove = doc2.paragraphs[second_ch4_para_idx:]
            for p in paras_to_remove:
                p._element.getparent().remove(p._element)
            print(f"用户文档：已删除从第{second_ch4_para_idx}段开始的重复章节内容")

    doc2.save(path_user)
    print(f"用户文档已保存：{path_user}")
    print("=== 所有修复完成 ===")

except ImportError:
    print("缺少 python-docx，正在安装...")
    import subprocess
    subprocess.run(['pip', 'install', 'python-docx'], check=True)
    print("安装完成，请重新运行此脚本")
except Exception as e:
    import traceback
    print(f"错误: {e}")
    traceback.print_exc()
