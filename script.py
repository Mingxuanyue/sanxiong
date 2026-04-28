# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(24)
    return p

def bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')

def tbl(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r in rows:
        row = t.add_row()
        for i, c in enumerate(r):
            row.cells[i].text = c
    doc.add_paragraph("")
    return t

# 封面
heading(doc, "三雄争锋", 0)
p = doc.add_paragraph("游戏用户文档  v2.0    2026年4月26日")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

# 第一章
heading(doc, "第一章  游戏简介", 1)
body(doc, "《三雄争锋》是一款三人纸牌对战游戏，在Windows命令行中运行。三名玩家各持18张牌，共18轮出牌比较大小积累分数，率先达到目标分者结算胜负。设计为三人实时对战，AI仅作无法联机时的替代方案。")

# 第二章
heading(doc, "第二章  快速开始", 1)

heading(doc, "2.1 运行游戏", 2)
bullet(doc, "编译环境：Dev-C++ 5.11，编译器选项添加 -std=c++11。")
bullet(doc, "文件编码：源文件使用 ANSI/GBK 编码，避免中文乱码。")
bullet(doc, "编译后双击运行生成的 .exe 文件，或在命令行中执行。")

heading(doc, "2.2 主菜单", 2)
tbl(doc,
    ['选项', '说明'],
    [
        ['1. 开始游戏', '进入模式选择'],
        ['2. 查看规则', '显示规则说明，按回车返回'],
        ['0. 退出游戏', '退出程序'],
    ]
)

heading(doc, "2.3 选择游戏模式", 2)
tbl(doc,
    ['模式', '说明'],
    [
        ['单局模式', '18轮结束按总分排名，并列算并列，不加赛'],
        ['达标分模式', '选定目标分（30/60/100/150），有人达标则结算或加赛'],
    ]
)

heading(doc, "2.4 选择初始规则", 2)
body(doc, "持有黑桃A的玩家选择本场初始规则（仅首场）：")
bullet(doc, "输入 1：比大——K最大，A最小，大王/小王均为最优")
bullet(doc, "输入 2：比小——A最小（最优），K最大（最差），大王/小王仍为最优")
body(doc, "达标分续场：自动继承当前规则，跳过选择；持有黑桃A者获得 +1分 补偿。")
body(doc, "AI持有黑桃A时：AI随机选择，界面公示结果，按回车继续。")

# 第三章
heading(doc, "第三章  每局操作指南", 1)

heading(doc, "3.1 正常出牌", 2)
body(doc, "每局开始时界面显示你的手牌：")
doc.add_paragraph("    [1] 黑桃K   [2] 红心5   [3] 梅花A   [4] 大王  ...")
body(doc, "输入方括号内的数字选择出哪张牌，按回车确认。输入无效时程序提示重新输入。")

heading(doc, "3.2 激活5号牌（狸猫换太子）", 2)
body(doc, "亮牌后，若满足以下全部条件，程序会询问是否激活：")
bullet(doc, "三张牌各不同")
bullet(doc, "你不是本轮最大（持有大王/小王者无法激活）")
bullet(doc, "手中有另一张5可消耗（本轮打出的5不算）")
bullet(doc, "本轮未触发出10，未触发连顺")
body(doc, "选择激活：消耗手中一张5，本轮正常计分，下轮开始前你本轮出的牌自动收回手中。")
body(doc, "注意：收回的牌下一轮打出时不会触发连顺。")
body(doc, "若单纯打出5但不激活：5作为普通牌参与大小比较，无任何效果。")

heading(doc, "3.3 特殊情况：你出了10（且唯一）", 2)
body(doc, "触发顺路拐带，本轮不计分，依次执行：")
bullet(doc, "第①步：你自动获得另外两人打出的牌，加入手中。")
bullet(doc, "第②步：程序提示你自选弃1张（只有你知道弃了什么）。")
bullet(doc, "第③步：你的顺时针下家对你的手牌盲输入编号，弃掉对应位置的牌（双方均不知道被弃了什么）。")
body(doc, "若同一轮三张牌也构成连顺：弃牌流程全部完成后，再执行连顺抽牌。")

heading(doc, "3.4 特殊情况：三张牌点数连续（连顺抽牌）", 2)
body(doc, "触发条件：三张牌点数连续（如3、4、5 或 J、Q、K），非第18局，且本轮无5号牌激活。")
body(doc, "抽牌规则：")
bullet(doc, "抽取人：看不见目标手牌，界面只显示对方牌数，输入编号盲选。")
bullet(doc, "被抽人：看得见自己被拿走了哪张牌，界面公示。")
bullet(doc, "旁观者：知道动作发生，不知道具体内容。")
body(doc, "若你是总分最高者，你先抽；按 玩家A→B→C→A 顺序依次进行；所有人抽完后统一加入手中。")

heading(doc, "3.5 特殊情况：你是被操作目标（使用大王）", 2)
body(doc, "当有人对你的手牌盲选编号N（抽牌或强制弃牌），且你手持大王、本场使用次数未耗尽时：")
bullet(doc, "程序提示：对方选了第N号，是否使用大王将目标偏移至第N-1号？")
bullet(doc, "选择使用：对方实际操作第N-1号位置的牌（1号偏移至末位）；大王不消耗，但使用次数-1。")
bullet(doc, "手牌只剩1张时：大王本身被选中，无法偏移。")
body(doc, "每人每场限用2次；续场时重置；加赛中禁用。")

heading(doc, "3.6 特殊情况：你是连顺抽牌的抽取人（使用小王）", 2)
body(doc, "若你手持小王、本场使用次数未耗尽，抽牌前程序会询问是否使用小王偷窥目标手牌：")
bullet(doc, "选择使用：短暂显示目标手牌内容，再由你输入编号盲选；小王不消耗，使用次数-1。")
bullet(doc, "每人每场限用1次；续场时重置；加赛中禁用。")

doc.save(r'd:\Users\ymx36\Desktop\三雄争锋游戏\游戏用户文档v2.docx')
print("Part3 OK")
