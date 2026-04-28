# -*- coding: utf-8 -*-
"""update_docs_v3b.py - 更新v3文档：随机AI + 5号牌说明修正"""
from docx import Document

def rep(p, old, new):
    full = p.text
    if old not in full:
        return False
    replaced = full.replace(old, new)
    if p.runs:
        for r in p.runs:
            r.text = ''
        p.runs[0].text = replaced
    else:
        p.add_run(replaced)
    return True

def rep_all(doc, old, new):
    count = 0
    for p in doc.paragraphs:
        if rep(p, old, new):
            count += 1
    return count

# ─── 用户文档 ─────────────────────────────────────────────────────────────
doc_u = Document('游戏用户文档v3.docx')
paras = doc_u.paragraphs

# 1. AI黑桃A → 智能选择规则
n = rep_all(doc_u,
    'AI持有黑桃A时：AI根据手牌智能选择（偏大牌则选比大，偏小牌则选比小），弹框公示结果，点击继续。',
    'AI持有黑桃A时：AI根据手牌智能选择（偏大牌则选比大，偏小牌则选比小），弹框公示结果，点击继续。')
# 可能尚未更新过，用原始文字
if n==0:
    rep_all(doc_u,
        'AI持有黑桃A时：AI随机选择，弹框公示结果，点击继续。',
        'AI持有黑桃A时：AI根据手牌智能选择（偏大牌则选比大，偏小牌则选比小），弹框公示结果，点击继续。')

# 2. 删除"收回的牌不触发连顺"说明（清空文字）
for p in paras:
    if '收回的牌下一轮打出时不会触发连顺' in p.text:
        for r in p.runs:
            r.text = ''
        print('  ✓ 用户文档 删除：收回牌连顺说明')

# 3. 5号牌FAQ改为正确表述
n = rep_all(doc_u,
    'Q：5号牌激活后，收回的牌下回合能组成连顺吗？',
    'Q：5号牌激活时，被消耗的那张5会参与连顺判定吗？')
print(f'  {"✓" if n else "✗"} 用户文档 5号牌FAQ问题 x{n}')

n = rep_all(doc_u,
    'A：不能。被5号牌收回的牌，在下一连合打出时不会触发连顺效果。',
    'A：不会。5号牌是消耗品，消耗的5本轮已出牌，不额外构成或破坏连顺；连顺判定基于三人本轮打出的实际牌。')
print(f'  {"✓" if n else "✗"} 用户文档 5号牌FAQ答案 x{n}')

doc_u.save('游戏用户文档v3_new.docx')
print('用户文档 v3 已保存为 游戏用户文档v3_new.docx（原文件被Word占用，请手动替换）\n')

# ─── 设计文档 ─────────────────────────────────────────────────────────────
doc_d = Document('游戏设计文档v3.docx')

# 1. AI黑桃A → 智能选择
n = rep_all(doc_d,
    'AI持有黑桃A时：随机选择规则，界面公示结果。',
    'AI持有黑桃A时：根据手牌智能选择（非Joker平均点数>7则选比大，否则选比小），界面公示结果。')
print(f'  {"✓" if n else "✗"} 设计文档 黑桃A规则选择 x{n}')

# 2. 5.1 标题 - 恢复为只代表玩家B
n = rep_all(doc_d,
    '5.1 策略AI（替代玩家B / C，v3起两个AI策略相同）',
    '5.1 策略AI（替代玩家B）')
print(f'  {"✓" if n else "✗"} 设计文档 5.1标题 x{n}')

# 3. 5.2 随机AI - 改为正确描述
n = rep_all(doc_d,
    '5.2 AI（替代玩家C）',
    '5.2 随机AI（替代玩家C）')
print(f'  {"✓" if n else "✗"} 设计文档 5.2标题 x{n}')

n = rep_all(doc_d,
    'v3起与策略AI（玩家B）采用相同的位置感知策略（strategicChoice），不再使用纯随机。',
    '一般回合出牌：完全随机选择（不考虑局面与牌的强弱），不可预测。')
print(f'  {"✓" if n else "✗"} 设计文档 5.2描述 x{n}')

n = rep_all(doc_d,
    '（策略与5.1相同，详见上方）',
    '特殊处理（大王防御、小王偷窥、出10弃牌、5号牌激活）：与5.1策略AI完全相同。')
print(f'  {"✓" if n else "✗"} 设计文档 5.2特殊处理说明 x{n}')

# 4. 删除设计文档中"收回的牌不构成连顺"类说明
for p in doc_d.paragraphs:
    if '收回的牌' in p.text and '连顺' in p.text:
        for r in p.runs:
            r.text = ''
        print('  ✓ 设计文档 删除：收回牌连顺说明')

import os as _os, shutil as _shutil
doc_d.save('游戏设计文档v3_new.docx')
print('设计文档 v3 已保存为 游戏设计文档v3_new.docx\n')
print('全部完成！')
